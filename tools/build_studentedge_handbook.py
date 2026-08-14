from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "StudentEdge_System_Handbook.docx"
LOGO = ROOT / "public" / "images" / "studentedge-mark.png"

INK = "241E1B"
MUTED = "786D66"
ACCENT = "A8567B"
ACCENT_DARK = "743951"
PALE = "F5E9EE"
SOFT = "F7F3F1"
LINE = "D8C5CC"
WHITE = "FFFFFF"


def font(run, size=None, bold=None, color=INK, name="Aptos", italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[i]))
            cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        shade(table.rows[0].cells[i], ACCENT_DARK)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(header), 9, True, WHITE)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if r_idx % 2:
                shade(cells[i], SOFT)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(str(value)), 9.2)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    font(p.add_run(text), 10.3)
    return p


def add_step(doc, title, detail):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(title + ". "), 10.5, True, ACCENT_DARK)
    font(p.add_run(detail), 10.5)


def callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, PALE)
    cell_margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label.upper() + "  "), 9, True, ACCENT_DARK)
    font(p.add_run(text), 10)
    set_table_widths(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def section_title(doc, number, title, subtitle):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    font(p.add_run(f"PART {number}"), 9, True, ACCENT)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    font(p.add_run(title), 22, True, ACCENT_DARK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    font(p.add_run(subtitle), 11.5, False, MUTED, italic=True)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.3)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18
for name, size, color, before, after in (
    ("Heading 1", 18, ACCENT_DARK, 16, 8),
    ("Heading 2", 14, ACCENT_DARK, 13, 6),
    ("Heading 3", 11.5, ACCENT, 10, 4),
):
    style = styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(hp.add_run("STUDENTEDGE  /  SYSTEM HANDBOOK"), 8, True, MUTED)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(fp.add_run("StudentEdge - Politeknik Besut  |  Internal reference  |  August 2026"), 8, False, MUTED)

# Cover
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(65)
    picture = p.add_run().add_picture(str(LOGO), width=Inches(0.9))
    doc_pr = picture._inline.docPr
    doc_pr.set("descr", "StudentEdge crest")
    doc_pr.set("title", "StudentEdge")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(8)
font(p.add_run("StudentEdge"), 30, True, ACCENT_DARK, "Aptos Display")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
font(p.add_run("System Guideline, Workflow, Coding & Function Handbook"), 17, True, INK, "Aptos Display")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(38)
font(p.add_run("A practical reference for users, administrators, lecturers, support staff, and developers"), 11, False, MUTED, italic=True)
add_table(doc, ["Document", "Current implementation baseline"], [
    ("System", "StudentEdge / e-Biasiswa Student Affairs Platform"),
    ("Edition", "August 2026"),
    ("Coverage", "Guidelines, operational workflows, coding architecture, and 185 named application functions"),
    ("Audience", "Students, lecturers, administrators, guards, JHEP staff, system administrators, and maintainers"),
], [2300, 7060])
callout(doc, "Document rule", "Use this handbook together with role permissions and institutional policy. Screens and functions may be hidden when a feature flag or role permission is disabled.")

# Contents
doc.add_page_break()
doc.add_heading("How to use this handbook", level=1)
doc.add_paragraph("The handbook deliberately follows the requested order: guidance first, operational workflows second, coding explanation third, and the implemented function catalogue last.")
add_table(doc, ["Part", "Purpose", "Primary reader"], [
    ("1. Guidelines", "Safe, consistent, privacy-aware use of the platform", "All users"),
    ("2. Workflows", "Step-by-step paths for daily operations", "Students and operational staff"),
    ("3. Coding explanation", "How requests, data, permissions, files, AI, and the PWA are implemented", "Developers and maintainers"),
    ("4. Function catalogue", "Role- and module-organized implemented capability inventory", "Project team, assessors, buyers, and auditors"),
    ("5. Operations", "Verification, deployment, backup, and troubleshooting", "System administrators and developers"),
], [1050, 5160, 3150])
doc.add_heading("Current verified implementation snapshot", level=2)
add_bullet(doc, "185 named application endpoints: 26 public/shared, 26 student, 7 lecturer AI Helper, and 126 administration endpoints.")
add_bullet(doc, "134 public controller methods across shared, authentication, student, and administration controllers.")
add_bullet(doc, "Two scheduled background operations: notification reminders and AI conversation retention cleanup.")
add_bullet(doc, "English and Malay interfaces, responsive PWA shell, light/dark themes, push notifications, and private authenticated document delivery.")

# Part 1 Guidelines
section_title(doc, "1", "Guidelines", "Rules for secure, accurate, and consistent use")
doc.add_heading("1.1 General user guidelines", level=2)
for item in [
    "Use only your own account. Never share passwords, QR sessions, reset codes, or authenticated devices.",
    "Confirm names, matric numbers, dates, amounts, statuses, and uploaded evidence before submitting a form.",
    "Upload readable source files in the supported format and remove unrelated personal information where possible.",
    "Treat AI responses as drafts. Verify important conclusions against the source document and the relevant module record.",
    "Log out from shared devices and revoke unknown active sessions from Settings.",
    "Report defects through Report a Problem with a clear description, page location, expected result, actual result, and screenshot.",
]: add_bullet(doc, item)
doc.add_heading("1.2 Privacy and authorization", level=2)
callout(doc, "Privacy", "Student profiles, identity numbers, payment evidence, offer letters, movement records, and private documents must be accessed only for an authorized institutional purpose.")
for item in [
    "Students may view only their own private records and documents.",
    "Lecturer and staff access is controlled by role, scope, and individually enabled pages.",
    "Sensitive student viewing, export, mutation, scholarship, discipline, movement, document, and system-administration permissions are intentionally separate.",
    "Private files must be served through authenticated controllers; do not expose private storage through a public URL.",
    "System Admin bulk deletion, public laptop identity workflows, and production data exports require explicit institutional approval and audit review.",
]: add_bullet(doc, item)
doc.add_heading("1.3 AI Helper guidelines", level=2)
add_table(doc, ["Role", "Expected AI behaviour", "Data boundary"], [
    ("Student", "Strict assistance for the student's own portal context", "No cross-student or administrative data"),
    ("Lecturer", "Natural conversation, research, document analysis, and permitted institutional context", "Only assigned lecturer scopes and enabled pages"),
    ("Administrator", "Research, summaries, reports, document analysis, and authorized system context", "Only abilities granted to the current account"),
    ("System administrator", "Full configured administrative AI capability", "Still subject to privacy, audit, and source-verification rules"),
], [1500, 4080, 3780])
for item in [
    "When a file is attached, ask the AI to analyze that file explicitly; attached PDF, image, CSV, and XLSX content must remain distinguishable from system context.",
    "Do not accept a generated report as an official decision. Use Edit or selected-text actions to refine it, then verify and approve through the real module.",
    "Conversation history is account-scoped and inactive conversations are automatically removed according to the configured retention period.",
    "Regeneration is queued/locked while a request is running to prevent duplicate generations.",
]: add_bullet(doc, item)
doc.add_heading("1.4 Data-entry and file standards", level=2)
add_table(doc, ["Item", "Guideline"], [
    ("Student import", "Use the supplied CSV/XLSX column structure; validate duplicates, program identifiers, matric number, and identity number before import."),
    ("Scholarship/B40", "Use the relevant import/export flow and review eligibility/status after processing."),
    ("Evidence", "Use clear JPG, JPEG, PNG, or WebP images within the configured limit."),
    ("AI source files", "Use PDF, supported images, CSV, or XLSX; keep filenames descriptive and avoid oversized or password-protected files."),
    ("Private documents", "Use the Student Document Centre; review status and expiry dates, and never place private documents under public storage."),
], [1900, 7460])

# Part 2 Workflows
section_title(doc, "2", "Operational workflows", "End-to-end paths from sign-in to completed action")
doc.add_heading("2.1 Authentication and account recovery", level=2)
for title, detail in [
    ("Open the login page", "Enter the registered identifier and password"),
    ("Establish the session", "The server validates the account, role, and active-device session"),
    ("Enter the correct portal", "Students, staff, guards, and administrators are routed to their authorized interface"),
    ("Recover when necessary", "Request a reset code, verify it, set a new password, and invalidate affected sessions"),
    ("Review devices", "Use Settings to revoke one device or all other sessions"),
]: add_step(doc, title, detail)
doc.add_heading("2.2 Student scholarship workflow", level=2)
for title, detail in [
    ("Review opportunities", "Open Scholarship and read current opportunities and announcements"),
    ("Open declaration", "Complete the student scholarship status declaration"),
    ("Attach supporting document", "Upload the offer letter when applicable"),
    ("Submit", "Confirm the information and submit the declaration"),
    ("Administrative review", "Authorized scholarship staff review the declaration and download the private offer letter"),
]: add_step(doc, title, detail)
doc.add_heading("2.3 Discipline, fine, and payment workflow", level=2)
for title, detail in [
    ("Record offense", "Authorized staff select the student, offense type, date, location, fine, and evidence"),
    ("Notify and review", "The student sees the offense and can print the record"),
    ("Submit payment evidence", "The student uploads a payment receipt for the relevant offense"),
    ("Decide application", "Authorized staff approve or reject the fine-payment application"),
    ("Update status", "The official offense/payment status is updated and remains auditable"),
]: add_step(doc, title, detail)
doc.add_heading("2.4 Vehicle sticker workflow", level=2)
for title, detail in [
    ("Student application", "Enter vehicle details and attach required evidence"),
    ("Administrative review", "Filter applications and inspect the submitted documents"),
    ("Decision", "Approve or reject the application with the official status"),
    ("Reporting", "Export application data when authorized"),
]: add_step(doc, title, detail)
doc.add_heading("2.5 Campus movement workflow", level=2)
for title, detail in [
    ("Open scanner", "The student opens the movement QR scanner"),
    ("Scan guard-house QR", "The active QR token is validated"),
    ("Checkout or return", "The system records the correct transition atomically"),
    ("Monitor", "Guards and authorized staff view outside students, records, deadlines, and violations"),
    ("Report", "Authorized staff filter and export movement records"),
]: add_step(doc, title, detail)
doc.add_heading("2.6 Student document workflow", level=2)
for title, detail in [
    ("Upload", "The student adds a private document and relevant metadata"),
    ("Review", "Authorized administrators inspect and approve or reject the submission"),
    ("Download", "The owner or authorized reviewer downloads through an authenticated controller"),
    ("Expiry reminder", "The scheduled reminder notifies subscribed students near configured expiry dates"),
]: add_step(doc, title, detail)
doc.add_heading("2.7 AI research and report workflow", level=2)
for title, detail in [
    ("Start or resume a chat", "Open AI Helper, choose New chat, or select an account-scoped history item"),
    ("Add sources", "Attach PDF, image, CSV, or XLSX files when the answer must be grounded in external material"),
    ("Send a clear request", "Ask a natural question and state whether the AI should use attached files, system data, or both"),
    ("Review source boundaries", "Confirm that file findings came from the attachment and system findings came from authorized records"),
    ("Refine", "Regenerate once, edit the full output, or select text and ask for focused changes"),
    ("Finalize outside AI", "Copy the reviewed draft into the official workflow and complete the real approval or publication process"),
]: add_step(doc, title, detail)
doc.add_heading("2.8 Laptop borrowing workflow", level=2)
for title, detail in [
    ("Prepare inventory", "Authorized staff register devices and print QR labels"),
    ("Identify borrower", "Use the public or staff scanning flow and validate the borrower"),
    ("Scan device", "The transaction processes an atomic borrow or return action"),
    ("Monitor", "Staff review current loans and transaction history"),
]: add_step(doc, title, detail)
doc.add_heading("2.9 Problem reporting workflow", level=2)
for title, detail in [
    ("Describe the issue", "Select a category and provide subject, expected behaviour, actual behaviour, and reproduction steps"),
    ("Attach a screenshot", "Include a clear image without unnecessary private information"),
    ("Submit", "The report is stored and delivery to the configured administrator is attempted"),
    ("Triage", "System administrators update status and notes"),
    ("Notify", "Push and email delivery status is recorded where configured"),
]: add_step(doc, title, detail)

# Part 3 Coding
section_title(doc, "3", "Coding explanation", "How the Laravel application is organized and how a request becomes a result")
doc.add_heading("3.1 Technology stack", level=2)
add_table(doc, ["Layer", "Implementation"], [
    ("Backend", "PHP 8.3+ and Laravel 13"),
    ("Database", "MySQL/MariaDB relational storage with migrations and imported baseline schema"),
    ("Frontend", "Blade views, Vite 8 assets, CSS design system, JavaScript interactions, and Tailwind CSS 4 tooling"),
    ("PWA", "Responsive application shell, service worker, install guidance, push subscription, and notification feed"),
    ("AI", "Provider abstraction with Gemini configuration, role-aware prompts, authorized system context, file extraction, and stored conversations"),
    ("Files", "Public assets plus private authenticated storage for sensitive student material"),
], [1850, 7510])
doc.add_heading("3.2 Request lifecycle", level=2)
for title, detail in [
    ("Route", "routes/web.php maps the HTTP method and URL to a controller or closure and assigns middleware and a route name"),
    ("Middleware", "Authentication, active-session checks, scope/ability checks, lecturer page gates, and feature flags reject unauthorized requests"),
    ("Controller", "The controller validates input, calls application services/helpers, reads or changes data, records audit information, and returns a view, redirect, file, or JSON response"),
    ("Data layer", "Query Builder and supporting services use the database schema and transactions where consistency matters"),
    ("Presentation", "Blade templates render server data; JavaScript adds chat, scanner, PWA, notification, and responsive interactions"),
    ("Verification", "Feature tests, PHP linting, route inspection, frontend builds, and UAT validate the result"),
]: add_step(doc, title, detail)
doc.add_heading("3.3 Main code map", level=2)
add_table(doc, ["Path", "Responsibility"], [
    ("routes/web.php", "Named public, shared, student, lecturer, and administration endpoints"),
    ("routes/console.php", "Scheduled reminders and AI-conversation retention cleanup"),
    ("app/Http/Controllers", "Request validation and orchestration grouped by Shared, Student, Auth, and Admin"),
    ("app/Services and app/Support", "Reusable session, permissions, AI, feature, settings, and operational logic"),
    ("resources/views", "Blade pages, reusable layouts, responsive components, and role navigation"),
    ("resources/css and resources/js", "Design tokens, PWA/mobile behaviour, chat UI, scanners, and client interactions"),
    ("database/migrations", "Incremental schema changes and indexes"),
    ("lang/en.json and lang/ms.json", "English/Malay display translations"),
    ("tests", "Regression and authorization coverage"),
    ("docs and Document Context", "System, architecture, operations, UAT, and product references"),
], [2700, 6660])
doc.add_heading("3.4 Authentication and authorization", level=2)
doc.add_paragraph("StudentEdge uses role-aware sessions and centralized ability checks. Route middleware provides the first gate; controllers must still enforce ownership and sensitive-data rules before reading, exporting, downloading, or mutating records.")
add_bullet(doc, "Authentication answers who the user is; authorization answers what the user is allowed to do.")
add_bullet(doc, "Role names do not replace granular abilities. Student list, sensitive detail, export, mutation, scholarship, discipline, movement, document, staff, and system permissions remain separable.")
add_bullet(doc, "Lecturer access combines account role, category/scope, enabled-page configuration, and global feature flags.")
add_bullet(doc, "Private downloads must check both authentication and ownership/ability before streaming a file.")
doc.add_heading("3.5 Database and consistency", level=2)
add_bullet(doc, "The imported StudentEdge.sql baseline and incremental Laravel migrations together define the current database.")
add_bullet(doc, "Use transactions for paired or stateful operations such as QR borrow/return and movement transitions.")
add_bullet(doc, "Large feeds use ordered cursor pagination, JSON batches, near-bottom loading, and DOM virtualization.")
add_bullet(doc, "Indexes support movement feeds, student profile lookup, and name-based student lists.")
add_bullet(doc, "Audit logs record sensitive administrative actions; production troubleshooting should preserve this evidence.")
doc.add_heading("3.6 AI Helper implementation", level=2)
add_table(doc, ["Component", "Behaviour"], [
    ("Role entry points", "Separate student, lecturer, and administrator routes apply different prompts, scopes, and feature gates."),
    ("Context", "The server builds only the system context authorized for the current account."),
    ("Attachments", "PDF/image/CSV/XLSX extraction is sent as a source context distinct from live system data."),
    ("Conversation storage", "Messages and titles are stored per account, can be reopened/renamed/deleted, and expire after inactivity."),
    ("Editing", "Authorized admin/lecturer users can edit generated messages; selected text can be sent back for focused revision."),
    ("Concurrency", "The interface disables/queues regeneration while a request is active to limit duplicate output."),
], [2200, 7160])
doc.add_heading("3.7 File and upload security", level=2)
add_bullet(doc, "Validate extension, MIME type, size, and ownership on the server; browser accept attributes are only a convenience.")
add_bullet(doc, "Generate storage filenames rather than trusting the supplied filename.")
add_bullet(doc, "Keep sensitive files in private storage and stream them only after authorization.")
add_bullet(doc, "Never cache authenticated pages or private downloads in the PWA service worker.")
add_bullet(doc, "Treat spreadsheet cells and extracted document text as untrusted input when displayed or used in prompts.")
doc.add_heading("3.8 Adding a new function safely", level=2)
for title, detail in [
    ("Define the requirement", "Identify the owner, permitted roles, data, validation, success state, and failure states"),
    ("Add authorization", "Create or reuse the smallest relevant ability and apply middleware plus controller checks"),
    ("Implement data changes", "Add a migration/index if required and preserve existing records"),
    ("Create the route and controller", "Use a named route, explicit validation, and auditable operations"),
    ("Build the interface", "Support responsive layouts, light/dark themes, Malay/English labels, keyboard access, and reduced motion"),
    ("Test", "Cover allowed access, forbidden access, validation, ownership, success, failure, and relevant performance paths"),
    ("Document and deploy", "Update function/system documentation, build assets, migrate safely, clear caches, and perform UAT"),
]: add_step(doc, title, detail)

# Part 4 Functions
section_title(doc, "4", "Implemented function catalogue", "A role- and module-organized view of the current system")
doc.add_heading("4.1 Count and method", level=2)
doc.add_paragraph("The reproducible total is 185 named application endpoints. It excludes the unnamed framework health route and two framework storage routes. Internal helper methods, UI-only controls, and scheduled commands are described separately so the number is not artificially inflated.")
add_table(doc, ["Area", "Count"], [
    ("Public and shared", "26"),
    ("Student", "26"),
    ("Lecturer AI Helper", "7"),
    ("Administration", "126"),
    ("Total", "185"),
], [6500, 2860])
doc.add_heading("4.2 Public and shared functions", level=2)
add_table(doc, ["Module", "Implemented capabilities"], [
    ("Home and overview", "Public home page and live system overview"),
    ("Authentication", "Login, logout, forgot password, code verification, and password reset"),
    ("Preferences", "Locale, theme, account settings, and linked role-mode switching"),
    ("Sessions", "View active devices; revoke one or all other sessions"),
    ("Notifications", "Notification feed plus Web Push subscribe/unsubscribe"),
    ("Support", "Public and authenticated problem-report submission"),
    ("Laptop public flow", "Open QR borrowing form, check borrower identity, and submit borrowing/return action"),
], [2100, 7260])
doc.add_heading("4.3 Student functions", level=2)
add_table(doc, ["Module", "Count", "Implemented capabilities"], [
    ("Dashboard", "1", "Operational overview and role-specific notices"),
    ("Profile", "3", "View/update profile and change password"),
    ("Scholarship", "4", "Records, announcements, declaration, and offer-letter submission"),
    ("Discipline", "5", "Offenses, print view, fine receipt, rules, and announcements"),
    ("Vehicle sticker", "2", "List and submit application with evidence"),
    ("Movement", "3", "History, scanner, and checkout/return submission"),
    ("Documents", "4", "List, upload, authenticated download, and delete own document"),
    ("AI Helper", "6", "Chat, history loading, rename, individual deletion, and delete all"),
], [1800, 780, 6780])
doc.add_heading("4.4 Lecturer functions", level=2)
doc.add_paragraph("Every lecturer receives only the functions allowed by the account's category, scope, enabled pages, and global feature flags. The dedicated AI Helper contributes seven named endpoints.")
add_table(doc, ["Area", "Capabilities"], [
    ("AI Helper", "Open helper, ask, load conversation, rename, edit AI message, delete one, and delete all"),
    ("Scholarship", "Assigned scholarship pages and student lookup where enabled"),
    ("Discipline", "Assigned offense, rules, fine, announcement, or sticker pages where enabled"),
    ("Movement", "Assigned movement monitoring and reporting pages where enabled"),
    ("Student data", "Only the permitted lookup/list/detail/export ability; sensitive detail is separate"),
], [1900, 7460])
doc.add_heading("4.5 Administration functions", level=2)
admin_rows = [
    ("Dashboard and monitoring", "2", "Dashboard analytics and live system monitoring"),
    ("Admin users", "7", "List/create/store/edit/update/reset password/delete"),
    ("Staff accounts", "8", "CRUD, import, reset password, and lecturer page configuration"),
    ("Guard accounts", "7", "CRUD and reset password"),
    ("Student management", "12", "List/search/detail/CRUD/import/export/reset/bulk delete"),
    ("Scholarships", "10", "CRUD/export plus B40 TVET view/import/export"),
    ("Scholarship announcements", "7", "CRUD and export"),
    ("Student scholarship status", "2", "Review declarations and download offer letters"),
    ("Offenses", "10", "List/create/store/export/edit/update/paid/delete/PDF/print"),
    ("Rules", "7", "CRUD and export"),
    ("Discipline announcements", "7", "CRUD and export"),
    ("Fine applications", "3", "List/export/approve or reject"),
    ("Vehicle stickers", "4", "List/export/decision/delete"),
    ("Movement", "10", "Records/export/outside/violations/QR/status/settings"),
    ("Student documents", "3", "Review list/authenticated download/decision"),
    ("Feature and session control", "3", "Feature flags and configured session lifetime"),
    ("Laptop management", "4", "Inventory/history, printable labels, scanner, and scan processing"),
    ("Maintenance", "5", "Control/cache state, push test, email test, maintenance broadcast"),
    ("Monthly report", "1", "Selected-month operational analytics report"),
    ("AI Helper", "7", "Chat/history/edit/rename/delete functions"),
    ("Bug reports", "3", "List/update/delete"),
    ("Admin profile", "3", "Profile, photo crop/upload, and password change"),
    ("Active visitors", "1", "System Admin active-session monitoring"),
]
add_table(doc, ["Module", "Count", "Implemented capabilities"], admin_rows, [2200, 700, 6460])
doc.add_heading("4.6 Scheduled functions", level=2)
add_table(doc, ["Schedule", "Function"], [
    ("Every five minutes", "Send subscribed students document-expiry and approaching movement-return reminders without overlapping runs"),
    ("Daily at 02:15", "Delete inactive student/admin AI conversations older than the configured retention period"),
], [2100, 7260])

# Part 5 Operations
section_title(doc, "5", "Operations and verification", "How to validate, deploy, recover, and troubleshoot the platform")
doc.add_heading("5.1 Local setup", level=2)
for title, detail in [
    ("Install dependencies", "Run composer install and npm install"),
    ("Configure environment", "Copy .env.example, set database/mail/push/AI values, and generate the application key"),
    ("Prepare database", "Import StudentEdge.sql, then run php artisan migrate"),
    ("Prepare storage", "Run php artisan storage:link only for intended public assets; keep private student files outside public storage"),
    ("Build", "Run npm run build"),
    ("Start", "Run php artisan serve or use the configured Laragon virtual host"),
]: add_step(doc, title, detail)
doc.add_heading("5.2 Verification checklist", level=2)
add_table(doc, ["Check", "Command or evidence"], [
    ("PHP syntax", "php -l <changed PHP file>"),
    ("Routes", "php artisan route:list --json"),
    ("Tests", "php artisan test --compact"),
    ("Frontend", "npm run build"),
    ("Formatting", "git diff --check"),
    ("Database", "Focused feature tests against a prepared test database"),
    ("UI", "Desktop and phone UAT in light/dark themes, keyboard navigation, reduced motion, and PWA mode"),
    ("Permissions", "Positive and negative role tests for every sensitive route"),
], [1900, 7460])
doc.add_heading("5.3 Backup and restore", level=2)
add_bullet(doc, "Back up the database, .env/configuration secrets through an approved secure process, and private uploads.")
add_bullet(doc, "Record application version, migration state, backup time, and responsible operator.")
add_bullet(doc, "Perform restoration in an isolated environment and verify counts, authentication, private downloads, core workflows, and audit evidence.")
add_bullet(doc, "Do not consider a backup complete until a restoration rehearsal succeeds.")
doc.add_heading("5.4 Troubleshooting sequence", level=2)
for title, detail in [
    ("Reproduce", "Record account role, page, action, expected result, actual result, timestamp, and browser/device"),
    ("Check authorization", "Confirm feature flags, role scope, lecturer page gate, ownership, and active session"),
    ("Check logs", "Review Laravel, browser console, queue/scheduler, mail, and push-delivery evidence"),
    ("Check data", "Confirm tables, migrations, related rows, storage paths, and database availability"),
    ("Isolate", "Test the smallest controller/route/service path and avoid changing unrelated modules"),
    ("Verify fix", "Run focused tests, full feasible regression checks, build assets, and repeat the original UAT path"),
]: add_step(doc, title, detail)
doc.add_heading("5.5 Known production cautions", level=2)
callout(doc, "Production gate", "Before institutional deployment, remove temporary password fallbacks, rotate secrets, disable debug mode, review dependency advisories, verify private-cache behaviour, rehearse backup/restore, and obtain approval for high-risk workflows.")
for item in [
    "The migrations do not recreate the complete original schema without the baseline SQL import.",
    "AI output can be incorrect or overconfident; important values and decisions require human verification.",
    "Receipt images and QR/OCR evidence are not proof of authentic payment without an official verification source.",
    "Public identity-number workflows and bulk deletion require throttling, logging, backups, and explicit policy approval.",
    "Database-dependent tests may fail when local MySQL is stopped even if PHP linting and frontend builds pass.",
]: add_bullet(doc, item)

# Appendix
doc.add_page_break()
doc.add_heading("Appendix A - Role responsibility summary", level=1)
add_table(doc, ["Role", "Primary responsibility", "Important boundary"], [
    ("Student", "Maintain own profile, applications, documents, movement, payments, and student AI conversations", "Own records only"),
    ("Lecturer/JHEP staff", "Perform assigned scholarship, discipline, movement, research, or administrative tasks", "Only enabled scopes/pages"),
    ("Guard", "Operate movement and permitted student lookup workflows", "No broad sensitive administration"),
    ("Scholarship administrator", "Manage scholarships, announcements, declarations, and related reporting", "Scholarship scope"),
    ("Discipline administrator", "Manage offenses, fines, rules, stickers, and related reporting", "Discipline scope"),
    ("Head of Student Affairs", "Cross-domain operational oversight", "No unrestricted system administration"),
    ("System administrator", "Accounts, features, monitoring, maintenance, delivery tests, and all registered abilities", "Institutional policy and audit still apply"),
], [1900, 4380, 3080])
doc.add_heading("Appendix B - Source-of-truth files", level=1)
for item in [
    "routes/web.php and routes/console.php for exposed and scheduled operations",
    "app/Http/Controllers for request orchestration",
    "database/migrations plus StudentEdge.sql for schema state",
    "app/Services, app/Support, middleware, and authorization helpers for shared rules",
    "resources/views, resources/css, and resources/js for the interface",
    "lang/en.json and lang/ms.json for display localization",
    "tests for executable regression expectations",
    "docs/SYSTEM_DOCUMENTATION.md, docs/UAT_CHECKLIST.md, docs/BACKUP_RESTORE_SOP.md, and Document Context files for maintained project guidance",
]: add_bullet(doc, item)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("END OF HANDBOOK"), 9, True, ACCENT)

doc.core_properties.title = "StudentEdge System Guideline, Workflow, Coding & Function Handbook"
doc.core_properties.subject = "Operational and technical handbook for StudentEdge"
doc.core_properties.author = "StudentEdge Project Team"
doc.core_properties.keywords = "StudentEdge, guideline, workflow, coding, functions, Laravel, student affairs"
doc.core_properties.comments = "Generated from the implementation baseline on 12 August 2026."
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
