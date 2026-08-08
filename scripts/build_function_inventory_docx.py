from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "StudentEdge_Function_Inventory.docx"
NAVY = "18324A"
TEAL = "147D80"
GOLD = "C6923F"
PALE = "EAF4F4"
LIGHT = "F4F6F8"
TEXT = RGBColor(36, 43, 48)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        shade(cell, NAVY)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = Inches(widths[i])
            shade(cells[i], "FFFFFF" if row_index % 2 == 0 else LIGHT)
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(8.5)
            run.font.color.rgb = TEXT
            if i == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(.72)
section.bottom_margin = Inches(.72)
section.left_margin = Inches(.8)
section.right_margin = Inches(.8)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].font.color.rgb = TEXT
styles["Normal"].paragraph_format.space_after = Pt(5)
for name, size, color in (("Title", 28, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 12, TEAL)):
    style = styles[name]
    style.font.name = "Aptos Display"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(10)
    style.paragraph_format.space_after = Pt(5)

title = doc.add_paragraph(style="Title")
title.add_run("StudentEdge\nFunction Inventory")
subtitle = doc.add_paragraph()
subtitle.add_run("Implementation-audited reference | 8 August 2026 | Release 5d5ec84").italic = True
subtitle.paragraph_format.space_after = Pt(14)

callout = doc.add_table(rows=1, cols=1)
callout.alignment = WD_TABLE_ALIGNMENT.CENTER
set_repeat_header(callout.rows[0])
cell = callout.cell(0, 0)
shade(cell, PALE)
margins(cell, 180, 220, 180, 220)
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("142 implemented named application functions")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor.from_string(TEAL)
p.add_run("\n23 public/shared  |  21 student  |  98 admin")

doc.add_heading("Counting method", level=1)
doc.add_paragraph(
    "Each named application endpoint returned by Laravel route inspection is counted once. "
    "Laravel vendor health/storage routes, internal helper methods, scheduled commands, and UI-only controls are excluded. "
    "This gives the team a reproducible count of callable system functions without inflating the number."
)
add_table(doc, ["Area", "Count", "Share"], [
    ("Public and shared", 23, "16.2%"),
    ("Student", 21, "14.8%"),
    ("Admin", 98, "69.0%"),
    ("Total", 142, "100%"),
], [3.7, 1.0, 1.4])

doc.add_heading("Public and shared functions", level=1)
add_table(doc, ["Module", "Count", "Functions"], [
    ("Home and overview", 2, "Public home; live overview metrics"),
    ("Authentication and recovery", 9, "Login/logout; forgot, verify, and reset password forms/actions"),
    ("Preferences, roles, sessions", 8, "Locale; theme; settings; linked role mode; device/session revocation"),
    ("Notifications and support", 4, "Notification feed; push subscribe/unsubscribe; problem report"),
], [2.0, .7, 3.4])

doc.add_heading("Student functions", level=1)
student_rows = [
    ("Dashboard", 1, "Operational student summary"),
    ("Profile", 3, "View/update profile; change password"),
    ("Scholarships", 4, "Records; announcements; declaration form/submission and offer letter"),
    ("Discipline", 5, "Offenses; print; fine receipt; rules; announcements"),
    ("Vehicle stickers", 2, "View and submit application"),
    ("Movement", 3, "History; QR scan; checkout/return"),
    ("Documents", 2, "Owned list and authenticated download"),
    ("AI Helper", 1, "Documented unavailable entry point"),
]
add_table(doc, ["Module", "Count", "Functions"], student_rows, [1.65, .65, 3.8])

doc.add_section(WD_SECTION.NEW_PAGE)
doc.add_heading("Admin functions", level=1)
admin_rows = [
    ("Dashboard and monitoring", 2, "Dashboard; live monitoring"),
    ("Admin users / lecturer access", 7, "CRUD; reset password; configure lecturer pages"),
    ("Students", 11, "CRUD; AJAX lookup; sensitive view; import/export; reset password"),
    ("Scholarships", 10, "CRUD/export; B40 TVET view/import/export"),
    ("Scholarship announcements", 7, "CRUD and export"),
    ("Scholarship declarations", 2, "Review list; offer-letter download"),
    ("Offenses", 10, "List; AJAX register; CRUD; export; print/PDF; mark paid"),
    ("Rules", 7, "CRUD and export"),
    ("Discipline announcements", 7, "CRUD and export"),
    ("Fine applications", 3, "List; export; approve/reject"),
    ("Vehicle stickers", 4, "List; export; approve/reject; delete"),
    ("Movement", 10, "Records/export; outside; violations; QR controls; settings"),
    ("Student documents", 3, "Review list; private download; approve/reject"),
    ("Features and sessions", 3, "Feature flags; configured session lifetime"),
    ("Maintenance and push", 4, "Maintenance/cache; private test; maintenance broadcast"),
    ("Monthly report", 1, "Selected-month operational analytics"),
    ("AI Helper", 2, "Open helper; submit read-only question"),
    ("Bug reports", 3, "List; update; delete"),
    ("Admin profile", 2, "View; crop/upload photo"),
]
add_table(doc, ["Module", "Count", "Functions"], admin_rows, [1.8, .6, 3.7])

doc.add_heading("Role coverage", level=1)
add_bullets(doc, [
    "Lecturer: limited AJAX student lookup and only offense pages enabled by System Admin.",
    "Guard: movement and limited non-sensitive student list/search.",
    "Scholarship Admin: scholarship workflows and permitted student list/search.",
    "Discipline Admin: discipline, movement, and permitted student management.",
    "Head of Student Affairs: scholarship, discipline, movement, students, exports, and documents; no system controls.",
    "System Admin: all registered abilities plus administration, monitoring, settings, maintenance, and push controls.",
])

doc.add_heading("Operational notes", level=1)
add_bullets(doc, [
    "Maintenance push broadcasts are sent immediately and describe the selected future maintenance window.",
    "Sending a maintenance announcement does not automatically enable maintenance mode.",
    "Student AI remains unavailable; it is counted because the implemented route returns that explicit state.",
    "Re-run the route-based audit whenever named routes are added, renamed, or removed.",
])

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("StudentEdge | Implemented Function Inventory | Internal Reference")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(110, 120, 128)

doc.core_properties.title = "StudentEdge Implemented Function Inventory"
doc.core_properties.subject = "Role and module organized inventory of 142 implemented named application functions"
doc.core_properties.author = "StudentEdge Team"
doc.save(OUTPUT)
print(OUTPUT)
